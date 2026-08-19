from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\Users\HUUTHANH\Downloads\SOP RootRotation Livestream App - updated.docx")


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Title", 22, RGBColor(11, 84, 92)),
        ("Heading 1", 16, RGBColor(11, 84, 92)),
        ("Heading 2", 13, RGBColor(30, 41, 59)),
        ("Heading 3", 11.5, RGBColor(30, 41, 59)),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("STANDARD OPERATING PROCEDURE\nROOT ROTATION LIVESTREAM")
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 84, 92)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Livestream Operations and HR Data Management")
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(100, 116, 139)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Version", "2.1"),
        ("Updated on", "19/08/2026"),
        ("System", "Root Rotation Livestream"),
        ("Architecture", "Google Sheets API + Website API + MongoDB"),
    ]
    for idx, (left, right) in enumerate(rows):
        meta.cell(idx, 0).text = left
        meta.cell(idx, 1).text = right
        shade(meta.cell(idx, 0), "D8F0EA")

    doc.add_paragraph("")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run(
        "This SOP lists the active spreadsheets, tabs, and data destinations used by recruitment, scheduling, payroll, review, and contracts."
    )
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    r.font.italic = True
    r.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_section(WD_SECTION.NEW_PAGE)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_sheet_map_table(doc: Document) -> None:
    rows = [
        ["Workbook", "Tab", "Role", "System status", "Used for", "Write-back / Save target"],
        ["HR_STREAMING_ MASTER FILE.xlsx", "Portfolio_Master", "Master schema", "Active", "Streamer / Host master data", "Read as authoritative master reference"],
        ["HR_STREAMING_ MASTER FILE.xlsx", "Support_Master", "Master schema", "Active", "Support staff master data", "Read as authoritative master reference"],
        ["HR_STREAMING_ MASTER FILE.xlsx", "Live_Session_Master", "Master schema", "Active", "Session scheduling and assignment", "Read / reconcile as scheduling master"],
        ["HR_STREAMING_ MASTER FILE.xlsx", "TikTok_Sales_Import", "Operational source", "Active", "TikTok raw sales and live report import", "Read for payroll and session evaluation"],
        ["HR_STREAMING_ MASTER FILE.xlsx", "Base_Salary_Card", "Rule source", "Active", "Hourly rate and commission logic", "Read for payroll and host grade logic"],
        ["HR_STREAMING_ MASTER FILE.xlsx", "Grade_Review", "Review output", "Active", "Host review and grading updates", "Write review output here when sync is approved"],
        ["HR_STREAMING_ MASTER FILE.xlsx", "Post_Live_Report", "Operational source", "Active", "Post-live metrics and quality signals", "Read for session evaluation"],
        ["HR_STREAMING_ MASTER FILE.xlsx", "Payroll_Sheet", "Finance output", "Active", "Final payroll export sheet", "Write payroll export here"],
        ["2. Lịch live và support live update.xlsx", "Thông tin Mẫu Live", "Recruitment intake", "Active project sheet", "Host application / recruitment profile intake", "Write recruitment sync and contract code updates here when website pushes back"],
        ["2. Lịch live và support live update.xlsx", "Thông tin Support Live", "Recruitment intake", "Active project sheet", "Support application / recruitment profile intake", "Write recruitment sync and contract code updates here when website pushes back"],
        ["2. Lịch live và support live update.xlsx", "Collect lịch live chính", "Availability mirror", "Transitional / compatibility", "Host availability collection mirror", "Write website availability sync here only when collect sync is used"],
        ["2. Lịch live và support live update.xlsx", "Collect lịch sp live", "Availability mirror", "Transitional / compatibility", "Support availability collection mirror", "Write website availability sync here only when collect sync is used"],
        ["2. Lịch live và support live update.xlsx", "LIVE STREAM SCHEDULE", "Legacy reference", "Do not use as master", "Old schedule view only", "No direct save target in the new architecture"],
        ["2. Lịch live và support live update.xlsx", "Lương + commission", "Reference sheet", "Reference only", "Host offer reference and compensation guidance", "No payroll final write-back target"],
        ["Livestream Payroll Workspace.xlsx", "Host", "Workspace", "Operational workspace", "Host payroll working sheet", "Used as workspace / manual check when needed"],
        ["Livestream Payroll Workspace.xlsx", "Support", "Workspace", "Operational workspace", "Support payroll working sheet", "Used as workspace / manual check when needed"],
        ["Livestream Payroll Workspace.xlsx", "Bảng lươngT7", "Workspace", "Operational workspace", "Weekly payroll working view", "Used as workspace / manual check when needed"],
        ["Livestream Payroll Workspace.xlsx", "Payroll_2026-08-03 and Payroll_2026-08-10", "Export tabs", "Generated output", "Historical payroll export tabs", "Written when payroll export is generated for that week"],
    ]

    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    for i, text in enumerate(rows[0]):
        header[i].text = text
        shade(header[i], "0F766E")
        for paragraph in header[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    for row_values in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = value
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
        if row_values[3] == "Active":
            shade(cells[3], "DCFCE7")
        elif "Do not use" in row_values[3]:
            shade(cells[3], "FEE2E2")
        elif "Reference" in row_values[3] or "Transitional" in row_values[3]:
            shade(cells[3], "FEF3C7")
        else:
            shade(cells[3], "E0F2FE")


def build_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    style_doc(doc)
    add_cover(doc)

    add_heading(doc, "1. Purpose", 1)
    doc.add_paragraph(
        "This SOP defines where operations data is stored, which tabs are active, and which sheets must be used by HR, operations, payroll, review, and contract workflows under the current system architecture."
    )

    add_heading(doc, "2. Source of Truth and Architecture Rule", 1)
    add_bullets(doc, [
        "The old master file and its legacy data structures are fully deprecated.",
        "The approved master structure is now centered on three primary schemas: Portfolio_Master, Support_Master, and Live_Session_Master.",
        "The application connects to the approved spreadsheet programmatically through the Google Sheets API.",
        "The older Shared Brain upload method is not the system connection method for this approved file.",
    ])

    add_heading(doc, "3. Data Storage Map by Workbook and Tab", 1)
    doc.add_paragraph(
        "Use the following table to identify exactly where each data domain lives and whether the system reads from it, writes back to it, or treats it as reference-only."
    )
    add_sheet_map_table(doc)

    add_heading(doc, "4. Rules for Each Project Workbook", 1)

    add_heading(doc, "4.1 HR_STREAMING_ MASTER FILE.xlsx", 2)
    add_bullets(doc, [
        "This workbook is the approved master workbook for livestream operations and HR data management.",
        "Portfolio_Master is the authoritative source for Host / Streamer identity and capability.",
        "Support_Master is the authoritative source for Support identity and level.",
        "Live_Session_Master is the authoritative scheduling schema for session-level planning and assignment.",
        "TikTok_Sales_Import, Base_Salary_Card, Grade_Review, Post_Live_Report, and Payroll_Sheet are active downstream operational tabs tied to payroll, review, and session analysis.",
    ])

    add_heading(doc, "4.2 2. Lịch live và support live update.xlsx", 2)
    add_bullets(doc, [
        "This workbook is still related to the project, but it is not the approved master scheduling workbook.",
        "Thông tin Mẫu Live stores Host recruitment intake data and can receive website sync-back updates such as contract code or profile refreshes.",
        "Thông tin Support Live stores Support recruitment intake data and can receive website sync-back updates such as contract code or profile refreshes.",
        "Collect lịch live chính and Collect lịch sp live are compatibility / mirror tabs for availability sync when the collect-sheet flow is still used.",
        "LIVE STREAM SCHEDULE is legacy reference only and must not be treated as the current master schedule source.",
        "Lương + commission is a compensation reference tab, not the final payroll master output.",
    ])

    add_heading(doc, "4.3 Livestream Payroll Workspace.xlsx", 2)
    add_bullets(doc, [
        "This workbook acts as a payroll workspace and review surface, not the operational master source.",
        "Host, Support, and Bảng lươngT7 are working tabs for payroll review or manual checking where needed.",
        "Payroll_YYYY-MM-DD tabs are generated export tabs used for weekly payroll snapshots.",
    ])

    add_heading(doc, "5. Save / Write-Back Rules", 1)
    add_numbered(doc, [
        "When the system updates master streamer data, the authoritative reference remains Portfolio_Master.",
        "When the system updates master support data, the authoritative reference remains Support_Master.",
        "When the system reconciles session assignment or scheduling references, the authoritative schema is Live_Session_Master.",
        "When payroll is exported, the final sheet target is Payroll_Sheet in HR_STREAMING_ MASTER FILE.xlsx unless a separate approved export procedure is documented.",
        "When recruitment data is synced back from website to sheet, Host data is written to Thông tin Mẫu Live and Support data is written to Thông tin Support Live.",
        "When availability mirror sync is used, Host availability is written to Collect lịch live chính and Support availability is written to Collect lịch sp live.",
    ])

    add_heading(doc, "6. Field Mapping Rule", 1)
    add_bullets(doc, [
        "Employee ID is the primary identity key across Portfolio_Master and Support_Master.",
        "Session ID is the primary identity key across Live_Session_Master and downstream schedule-related processes.",
        "Host_ID in session flows must resolve to Portfolio_Master.",
        "Support_ID in session flows must resolve to Support_Master.",
        "TikTok_Sales_Import is used as the raw session performance and payroll input source, but person validation must still resolve back to Portfolio_Master and Support_Master.",
    ])

    add_heading(doc, "7. Technical Note for IT and Operations", 1)
    add_bullets(doc, [
        "The system does not rely on manual workbook upload for the approved master file.",
        "The backend uses Google Sheets API integration to read and write approved spreadsheet ranges.",
        "Any SOP, training guide, or handoff document must reference the workbook and tab names above exactly as written.",
        "Any tab marked legacy, transitional, or reference-only must not be promoted into a master operational workflow without formal SOP revision.",
    ])

    add_heading(doc, "8. Final Governance Statement", 1)
    doc.add_paragraph(
        "If a workflow, dashboard, or person is unsure where data should be stored, this SOP section must be checked first. No team member should guess the destination sheet or tab. The workbook and tab map above is the required operating reference."
    )

    return doc


def main() -> None:
    doc = build_doc()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
