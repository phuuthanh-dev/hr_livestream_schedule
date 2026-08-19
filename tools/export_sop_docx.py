from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "system-sop-root-rotation-livestream.md"
OUTPUT = ROOT / "docs" / "RootRotation_Livestream_SOP_v3.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color: str = "D8E4E8", size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def set_base_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Title", 24, RGBColor(12, 79, 84)),
        ("Heading 1", 16, RGBColor(12, 79, 84)),
        ("Heading 2", 13, RGBColor(25, 41, 51)),
        ("Heading 3", 11.5, RGBColor(25, 41, 51)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
    document.styles["List Bullet"].font.name = "Arial"
    document.styles["List Bullet"].font.size = Pt(10.5)
    document.styles["List Number"].font.name = "Arial"
    document.styles["List Number"].font.size = Pt(10.5)


INLINE_TOKEN_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
NUMBERED_RE = re.compile(r"^\d+\.\s+")


def add_inline_runs(paragraph, text: str) -> None:
    chunks = INLINE_TOKEN_RE.split(text)
    for chunk in chunks:
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
            continue
        if chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(24, 96, 88)
            continue
        paragraph.add_run(chunk)


def format_table_text(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.style = document.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("STANDARD OPERATING PROCEDURE\nROOT ROTATION LIVESTREAM")
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(12, 79, 84)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    r = subtitle.add_run("Website HR, tuyển dụng, lịch rảnh, xếp lịch, payroll và hợp đồng")
    r.font.name = "Arial"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(86, 104, 117)

    meta = document.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    meta.columns[0].width = Inches(2.1)
    meta.columns[1].width = Inches(4.9)
    labels = [
        ("Phiên bản", "1.0"),
        ("Ngày cập nhật", "19/08/2026"),
        ("Nguồn vận hành chính", "Website + MongoDB"),
    ]
    for idx, (label, value) in enumerate(labels):
        meta.cell(idx, 0).text = label
        meta.cell(idx, 1).text = value
        set_cell_shading(meta.cell(idx, 0), "D8EFE8")
    format_table_text(meta)

    document.add_paragraph("")
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(8)
    rr = note.add_run(
        "Tài liệu này mô tả quy trình thao tác chuẩn cho nhân sự và Admin/Manager trong hệ thống Root Rotation Livestream."
    )
    rr.font.name = "Arial"
    rr.font.size = Pt(10.5)
    rr.font.italic = True
    rr.font.color.rgb = RGBColor(86, 104, 117)

    document.add_section(WD_SECTION.NEW_PAGE)


def add_markdownish_body(document: Document, text: str) -> None:
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            document.add_paragraph("")
            continue

        if stripped == "---":
            document.add_paragraph("")
            continue

        if stripped.startswith("# "):
            p = document.add_paragraph(stripped[2:].strip(), style="Heading 1")
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(5)
            set_paragraph_bottom_border(p, "D8E4E8", "4")
            continue

        if stripped.startswith("## "):
            p = document.add_paragraph(stripped[3:].strip(), style="Heading 2")
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(3)
            continue

        if stripped.startswith("### "):
            p = document.add_paragraph(stripped[4:].strip(), style="Heading 3")
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            continue

        if NUMBERED_RE.match(stripped):
            p = document.add_paragraph(style="List Number")
            add_inline_runs(p, stripped)
            continue

        if stripped.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:].strip())
            continue

        if stripped.startswith("[🎥 VIDEO TUTORIAL CẦN QUAY:"):
            box = document.add_table(rows=1, cols=1)
            box.style = "Table Grid"
            cell = box.cell(0, 0)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_runs(paragraph, stripped)
            set_cell_shading(cell, "FFF4CC")
            format_table_text(box)
            continue

        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_inline_runs(p, stripped)


def main() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    set_base_style(document)
    add_cover(document)
    add_markdownish_body(document, SOURCE.read_text(encoding="utf-8"))

    for sec in document.sections:
        add_page_number(sec)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
